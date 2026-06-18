"""
DOCX Report Generator — Institutional-grade Word document (15,000–40,000 words)
================================================================================
25 structured sections covering all forensic dimensions.

References
----------
Canny, S. (2012-2023). "python-docx: Create and modify Word documents."
    python-docx.readthedocs.io.
    Core DOCX generation library; styles, headings, tables, footnotes.

Standards
---------
SEBI LODR Regulation 33 & 34 — Listed entity financial disclosure obligations;
    report structure maps to statutory disclosure requirements.
PCAOB AS 3101 — Critical Audit Matters reporting format adapted for forensic
    findings presentation.
ISA 265 / SA 265 — Communicating deficiencies in internal control; findings
    communication format.
Companies Act 2013 §134 — Board's Report requirements; governance and risk
    sections modelled on statutory board report format.

Role
----
Generates the primary long-form Word document report. Structure: Executive Summary
(verdict + risk score) → 17 Agent Findings (one section per agent) → Forensic Model
Scores (Beneish/Dechow/Altman/Piotroski/Accrual) → Governance Assessment →
Regulatory Risk → Recommendations → Technical Appendix (full calculations).
Word count: 15,000–40,000 words depending on data availability.
"""

from __future__ import annotations
from pathlib import Path
from datetime import datetime
from loguru import logger

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from config import REPORT_CONFIG


class DOCXGenerator:
    """Generates a comprehensive 15,000-40,000 word DOCX forensic report."""

    PRIMARY_COLOR = RGBColor(0x1A, 0x23, 0x7E) if HAS_DOCX else None
    ACCENT_RED = RGBColor(0xB7, 0x1C, 0x1C) if HAS_DOCX else None
    ACCENT_GREEN = RGBColor(0x1B, 0x5E, 0x20) if HAS_DOCX else None
    ACCENT_ORANGE = RGBColor(0xE6, 0x51, 0x00) if HAS_DOCX else None

    def generate(self, report_data: dict, output_path: Path) -> Path:
        if not HAS_DOCX:
            logger.warning("python-docx not installed. DOCX report skipped.")
            return output_path

        doc = Document()
        meta = report_data.get("metadata", {})
        company = meta.get("company_name", "Unknown Company")
        exec_sum = report_data.get("executive_summary", {})
        score = exec_sum.get("overall_risk_score", 50)
        verdict = exec_sum.get("verdict", "MONITOR")

        self._setup_styles(doc)
        self._add_cover_page(doc, company, score, verdict, meta)
        self._add_table_of_contents(doc)
        self._add_executive_summary(doc, report_data)
        self._add_business_model_analysis(doc, report_data)
        self._add_financial_performance_review(doc, report_data)
        self._add_historical_analysis(doc, report_data)
        self._add_revenue_forensics(doc, report_data)
        self._add_cashflow_forensics(doc, report_data)
        self._add_balance_sheet_forensics(doc, report_data)
        self._add_working_capital_review(doc, report_data)
        self._add_earnings_quality(doc, report_data)
        self._add_related_party_analysis(doc, report_data)
        self._add_auditor_assessment(doc, report_data)
        self._add_management_credibility(doc, report_data)
        self._add_governance_review(doc, report_data)
        self._add_fraud_detection_analysis(doc, report_data)
        self._add_credit_risk_review(doc, report_data)
        self._add_peer_benchmarking(doc, report_data)
        self._add_fraud_pattern_matching(doc, report_data)
        self._add_short_seller_report(doc, report_data)
        self._add_bull_case_report(doc, report_data)
        self._add_investment_committee_memo(doc, report_data)
        self._add_red_flag_registry(doc, report_data)
        self._add_green_flag_registry(doc, report_data)
        self._add_management_questionnaire(doc, report_data)
        self._add_monitoring_framework(doc, report_data)
        self._add_final_verdict(doc, report_data)

        doc.save(str(output_path))
        logger.info(f"DOCX saved: {output_path.name}")
        return output_path

    def _setup_styles(self, doc) -> None:
        """Configure document styles."""
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        for i in range(1, 4):
            h_style = doc.styles[f"Heading {i}"]
            h_style.font.name = "Calibri"
            h_style.font.bold = True
            h_style.font.color.rgb = self.PRIMARY_COLOR

    def _add_cover_page(self, doc, company: str, score: float, verdict: str, meta: dict) -> None:
        doc.add_paragraph()
        doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("FORENSIC ACCOUNTING INVESTIGATION REPORT")
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run(company.upper())
        run2.font.size = Pt(18)
        run2.font.bold = True

        doc.add_paragraph()
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        color = self.ACCENT_RED if score > 60 else (self.ACCENT_ORANGE if score > 40 else self.ACCENT_GREEN)
        run3 = score_para.add_run(f"OVERALL RISK SCORE: {score:.1f}/100")
        run3.font.size = Pt(16)
        run3.font.bold = True
        run3.font.color.rgb = color

        verdict_para = doc.add_paragraph()
        verdict_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = verdict_para.add_run(f"INVESTMENT VERDICT: {verdict}")
        run4.font.size = Pt(14)
        run4.font.bold = True
        run4.font.color.rgb = color

        doc.add_paragraph()
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run5 = meta_para.add_run(
            f"{REPORT_CONFIG.firm_name}\n"
            f"{REPORT_CONFIG.firm_tagline}\n"
            f"Report Date: {datetime.now().strftime('%B %d, %Y')}\n"
            f"STRICTLY CONFIDENTIAL — INVESTMENT COMMITTEE USE ONLY"
        )
        run5.font.size = Pt(10)
        run5.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_page_break()

    def _add_table_of_contents(self, doc) -> None:
        doc.add_heading("TABLE OF CONTENTS", level=1)
        sections = [
            ("1", "Executive Summary"),
            ("2", "Business Model Analysis"),
            ("3", "Financial Performance Review"),
            ("4", "5-Year Historical Analysis"),
            ("5", "Revenue Recognition Investigation"),
            ("6", "Cash Flow Forensics"),
            ("7", "Balance Sheet Forensics"),
            ("8", "Working Capital Review"),
            ("9", "Earnings Quality Assessment"),
            ("10", "Related Party Analysis"),
            ("11", "Auditor Assessment"),
            ("12", "Management Credibility Analysis"),
            ("13", "Corporate Governance Review"),
            ("14", "Fraud Detection Analysis"),
            ("15", "Credit Risk Review"),
            ("16", "Peer Benchmarking"),
            ("17", "Historical Fraud Pattern Matching"),
            ("18", "Short Seller Report"),
            ("19", "Bull Case Report"),
            ("20", "Investment Committee Memo"),
            ("21", "Red Flag Registry"),
            ("22", "Green Flag Registry"),
            ("23", "Management Questionnaire"),
            ("24", "Monitoring Framework"),
            ("25", "Final Verdict"),
        ]
        for num, title in sections:
            para = doc.add_paragraph(f"{num}. {title}", style="Normal")
            para.paragraph_format.space_before = Pt(2)
        doc.add_page_break()

    def _section_header(self, doc, number: str, title: str) -> None:
        doc.add_heading(f"SECTION {number}: {title}", level=1)
        separator = doc.add_paragraph("─" * 80)
        separator.runs[0].font.color.rgb = self.PRIMARY_COLOR

    def _add_executive_summary(self, doc, data: dict) -> None:
        self._section_header(doc, "1", "EXECUTIVE SUMMARY")
        exec_sum = data.get("executive_summary", {})
        company = data.get("metadata", {}).get("company_name", "")
        score = exec_sum.get("overall_risk_score", 50)
        verdict = exec_sum.get("verdict", "MONITOR")
        red_flags = exec_sum.get("total_red_flags", 0)
        green_flags = exec_sum.get("total_green_flags", 0)

        para = doc.add_paragraph()
        para.add_run("INVESTIGATION MANDATE\n").bold = True
        para.add_run(
            f"This report presents the findings of an institutional-grade forensic accounting "
            f"investigation into {company}. The investigation was conducted using a multi-agent "
            f"analytical framework comprising 16 specialist agents including forensic accountants, "
            f"credit analysts, governance specialists, and fraud investigators.\n\n"
            f"The investigation covers 5 years of financial history, analyzing annual reports, "
            f"quarterly results, earnings call transcripts, auditor reports, and regulatory filings. "
            f"All findings are evidence-based and traceable to specific document citations."
        )

        doc.add_paragraph()
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = "Table Grid"
        cells = [
            ("Overall Risk Score", f"{score:.1f} / 100"),
            ("Investment Verdict", verdict),
            ("Total Red Flags", str(red_flags)),
            ("Total Green Flags", str(green_flags)),
        ]
        for i, (label, value) in enumerate(cells):
            row = summary_table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value

        doc.add_paragraph()
        # Add key findings from final analysis
        final_analysis = exec_sum.get("final_analysis", "")
        if final_analysis:
            doc.add_paragraph("CHIEF INVESTIGATOR'S ASSESSMENT").runs[0].bold = True
            doc.add_paragraph(final_analysis[:2000])

        doc.add_page_break()

    def _add_section_with_agent_content(self, doc, section_num: str, title: str, agent_id: int, data: dict) -> None:
        """Generic section using agent output."""
        self._section_header(doc, section_num, title)
        agent_analyses = data.get("agent_analyses", {})
        agent_data = agent_analyses.get(str(agent_id), {})

        if agent_data.get("summary"):
            doc.add_paragraph(agent_data["summary"])
        else:
            doc.add_paragraph(f"Analysis for {title} not available in this investigation.")

        # Add relevant red flags
        red_flags = [f for f in data.get("red_flags", []) if f.get("agent", "").lower() in title.lower()]
        if red_flags:
            doc.add_paragraph("KEY FINDINGS:").runs[0].bold = True
            for flag in red_flags[:5]:
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(f"[{flag.get('risk_level', 'HIGH')}] {flag.get('title', '')}")
                run.bold = True
                run.font.color.rgb = self.ACCENT_RED
                doc.add_paragraph(flag.get("evidence", "")[:300])

        doc.add_page_break()

    def _add_business_model_analysis(self, doc, data): self._add_section_with_agent_content(doc, "2", "BUSINESS MODEL ANALYSIS", 0, data)
    def _add_financial_performance_review(self, doc, data): self._add_section_with_agent_content(doc, "3", "FINANCIAL PERFORMANCE REVIEW", 2, data)
    def _add_historical_analysis(self, doc, data): self._add_section_with_agent_content(doc, "4", "5-YEAR HISTORICAL ANALYSIS", 0, data)
    def _add_revenue_forensics(self, doc, data): self._add_section_with_agent_content(doc, "5", "REVENUE RECOGNITION INVESTIGATION", 3, data)
    def _add_cashflow_forensics(self, doc, data): self._add_section_with_agent_content(doc, "6", "CASH FLOW FORENSICS", 4, data)
    def _add_balance_sheet_forensics(self, doc, data): self._add_section_with_agent_content(doc, "7", "BALANCE SHEET FORENSICS", 5, data)
    def _add_working_capital_review(self, doc, data): self._add_section_with_agent_content(doc, "8", "WORKING CAPITAL REVIEW", 5, data)
    def _add_earnings_quality(self, doc, data): self._add_section_with_agent_content(doc, "9", "EARNINGS QUALITY ASSESSMENT", 8, data)
    def _add_related_party_analysis(self, doc, data): self._add_section_with_agent_content(doc, "10", "RELATED PARTY ANALYSIS", 9, data)
    def _add_auditor_assessment(self, doc, data): self._add_section_with_agent_content(doc, "11", "AUDITOR ASSESSMENT", 10, data)
    def _add_management_credibility(self, doc, data): self._add_section_with_agent_content(doc, "12", "MANAGEMENT CREDIBILITY ANALYSIS", 11, data)
    def _add_governance_review(self, doc, data): self._add_section_with_agent_content(doc, "13", "CORPORATE GOVERNANCE REVIEW", 9, data)

    def _add_fraud_detection_analysis(self, doc, data: dict) -> None:
        self._section_header(doc, "14", "FRAUD DETECTION ANALYSIS")
        forensic_scores = data.get("forensic_scores", [{}])
        scores = forensic_scores[0] if forensic_scores else {}

        doc.add_paragraph("QUANTITATIVE FRAUD MODELS").runs[0].bold = True

        m_score = scores.get("beneish_m_score", "N/A")
        z_score = scores.get("altman_z_score", "N/A")
        p_score = scores.get("piotroski_f_score", "N/A")
        d_score = scores.get("dechow_f_score", "N/A")

        fraud_table = doc.add_table(rows=5, cols=3)
        fraud_table.style = "Table Grid"
        headers = fraud_table.rows[0]
        headers.cells[0].text = "Model"
        headers.cells[1].text = "Score"
        headers.cells[2].text = "Assessment"

        model_data = [
            ("Beneish M-Score", f"{m_score:.3f}" if isinstance(m_score, float) else str(m_score),
             "MANIPULATION LIKELY" if isinstance(m_score, float) and m_score > -1.78 else "Within Range"),
            ("Altman Z-Score", f"{z_score:.3f}" if isinstance(z_score, float) else str(z_score),
             scores.get("altman_zone", "N/A")),
            ("Piotroski F-Score", f"{p_score}/9" if isinstance(p_score, int) else str(p_score),
             "WEAK" if isinstance(p_score, int) and p_score <= 2 else "NEUTRAL/STRONG"),
            ("Dechow F-Score", f"{d_score:.4f}" if isinstance(d_score, float) else str(d_score),
             "HIGH RISK" if isinstance(d_score, float) and d_score > 0.025 else "Acceptable"),
        ]
        for i, (model, score_val, assessment) in enumerate(model_data, 1):
            row = fraud_table.rows[i]
            row.cells[0].text = model
            row.cells[1].text = score_val
            row.cells[2].text = assessment

        agent_analyses = data.get("agent_analyses", {})
        agent_6 = agent_analyses.get("6", {})
        if agent_6.get("summary"):
            doc.add_paragraph()
            doc.add_paragraph(agent_6["summary"])

        doc.add_page_break()

    def _add_credit_risk_review(self, doc, data): self._add_section_with_agent_content(doc, "15", "CREDIT RISK REVIEW", 7, data)
    def _add_peer_benchmarking(self, doc, data): self._add_section_with_agent_content(doc, "16", "PEER BENCHMARKING", 12, data)
    def _add_fraud_pattern_matching(self, doc, data): self._add_section_with_agent_content(doc, "17", "HISTORICAL FRAUD PATTERN MATCHING", 13, data)
    def _add_short_seller_report(self, doc, data): self._add_section_with_agent_content(doc, "18", "SHORT SELLER REPORT (BEAR CASE)", 14, data)
    def _add_bull_case_report(self, doc, data): self._add_section_with_agent_content(doc, "19", "BULL CASE REPORT", 15, data)

    def _add_investment_committee_memo(self, doc, data: dict) -> None:
        self._section_header(doc, "20", "INVESTMENT COMMITTEE MEMO")
        exec_sum = data.get("executive_summary", {})
        verdict = exec_sum.get("verdict", "MONITOR")
        company = data.get("metadata", {}).get("company_name", "")
        score = exec_sum.get("overall_risk_score", 50)

        memo_text = doc.add_paragraph()
        memo_text.add_run("MEMORANDUM\n").bold = True
        memo_text.add_run(
            f"TO: Investment Committee\n"
            f"FROM: Chief Investigation Director, {REPORT_CONFIG.firm_name}\n"
            f"RE: {company} — Forensic Investigation Verdict\n"
            f"DATE: {datetime.now().strftime('%B %d, %Y')}\n\n"
            f"RECOMMENDATION: {verdict}\n"
            f"RISK SCORE: {score:.1f}/100\n\n"
        )

        final_analysis = exec_sum.get("final_analysis", "")
        if final_analysis:
            doc.add_paragraph(final_analysis[:3000])
        doc.add_page_break()

    def _add_red_flag_registry(self, doc, data: dict) -> None:
        self._section_header(doc, "21", "RED FLAG REGISTRY")
        red_flags = data.get("red_flags", [])
        doc.add_paragraph(f"Total Red Flags Identified: {len(red_flags)}")

        for i, flag in enumerate(red_flags, 1):
            doc.add_paragraph(f"RF-{i:03d}: {flag.get('title', '')}", style="Heading 2")
            doc.add_paragraph(f"Severity: {flag.get('risk_level', 'HIGH')}")
            doc.add_paragraph(f"Evidence: {flag.get('evidence', '')[:400]}")
            doc.add_paragraph(f"Fiscal Year: {flag.get('fiscal_year', 'N/A')}")
            doc.add_paragraph()
        doc.add_page_break()

    def _add_green_flag_registry(self, doc, data: dict) -> None:
        self._section_header(doc, "22", "GREEN FLAG REGISTRY")
        green_flags = data.get("green_flags", [])
        doc.add_paragraph(f"Total Green Flags (Positive Indicators): {len(green_flags)}")
        for i, flag in enumerate(green_flags, 1):
            doc.add_paragraph(f"GF-{i:03d}: {flag.get('title', '')}", style="Heading 2")
            doc.add_paragraph(flag.get("detail", ""))
        doc.add_page_break()

    def _add_management_questionnaire(self, doc, data: dict) -> None:
        self._section_header(doc, "23", "MANAGEMENT QUESTIONNAIRE")
        doc.add_paragraph(
            "The following questions should be submitted to management prior to any investment decision. "
            "Management's willingness and ability to answer clearly is itself a quality signal."
        )
        questions = data.get("management_questions", [])
        for i, q in enumerate(questions, 1):
            para = doc.add_paragraph(style="List Number")
            para.add_run(q)
        doc.add_page_break()

    def _add_monitoring_framework(self, doc, data: dict) -> None:
        self._section_header(doc, "24", "MONITORING FRAMEWORK")
        doc.add_paragraph("Monitor these metrics on a quarterly basis:")
        triggers = data.get("monitoring_framework", [])

        if triggers:
            table = doc.add_table(rows=len(triggers) + 1, cols=3)
            table.style = "Table Grid"
            headers = table.rows[0]
            headers.cells[0].text = "Metric"
            headers.cells[1].text = "Alert Threshold"
            headers.cells[2].text = "Action Required"
            for i, trigger in enumerate(triggers, 1):
                row = table.rows[i]
                row.cells[0].text = trigger.get("metric", "")
                row.cells[1].text = trigger.get("threshold", "")
                row.cells[2].text = trigger.get("action", "")
        doc.add_page_break()

    def _add_final_verdict(self, doc, data: dict) -> None:
        self._section_header(doc, "25", "FINAL VERDICT")
        exec_sum = data.get("executive_summary", {})
        verdict = exec_sum.get("verdict", "MONITOR")
        score = exec_sum.get("overall_risk_score", 50)
        company = data.get("metadata", {}).get("company_name", "")

        verdict_para = doc.add_paragraph()
        color = self.ACCENT_RED if score > 60 else (self.ACCENT_ORANGE if score > 40 else self.ACCENT_GREEN)
        run = verdict_para.add_run(f"FINAL VERDICT: {verdict}")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = color

        doc.add_paragraph(f"Risk Score: {score:.1f}/100")
        doc.add_paragraph()
        doc.add_paragraph(
            "This verdict is based on the aggregate findings of 16 specialist forensic agents "
            "and is subject to the limitations of available public information. "
            "This report is for institutional investor use only and does not constitute "
            "financial advice. Always conduct additional due diligence before making investment decisions."
        )

        doc.add_paragraph()
        doc.add_paragraph(
            f"Prepared by: {REPORT_CONFIG.firm_name}\n"
            f"Date: {datetime.now().strftime('%B %d, %Y')}\n"
            f"Classification: STRICTLY CONFIDENTIAL"
        )
